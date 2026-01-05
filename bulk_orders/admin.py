# bulk_orders/admin.py
from django.contrib import admin
from django.utils.html import format_html
from django.urls import reverse, path
from django.http import HttpResponse
from django.db.models import Count, Q, Prefetch
from django.shortcuts import redirect, get_object_or_404
from django.contrib import messages
from django.template.loader import render_to_string
from django.core.paginator import Paginator
from django.utils import timezone
from django.conf import settings
from io import BytesIO
import xlsxwriter
import logging
from .models import BulkOrderLink, OrderEntry, CouponCode
from .utils import generate_coupon_codes

logger = logging.getLogger(__name__)


class CouponCodeInline(admin.TabularInline):
    model = CouponCode
    extra = 0
    readonly_fields = ["code", "is_used", "created_at"]
    fields = ["code", "is_used", "created_at"]
    can_delete = False
    max_num = 0
    show_change_link = True

    def has_add_permission(self, request, obj=None):
        return False
    
    def get_queryset(self, request):
        """✅ FIX: Only show coupons for THIS bulk order"""
        qs = super().get_queryset(request)
        # The parent object (bulk_order) is automatically filtered
        return qs


class HasCouponFilter(admin.SimpleListFilter):
    title = "Coupon Status"
    parameter_name = "has_coupon"

    def lookups(self, request, model_admin):
        return (
            ("yes", "Has Coupon"),
            ("no", "No Coupon"),
        )

    def queryset(self, request, queryset):
        if self.value() == "yes":
            return queryset.filter(coupon_used__isnull=False)
        if self.value() == "no":
            return queryset.filter(coupon_used__isnull=True)


@admin.register(OrderEntry)
class OrderEntryAdmin(admin.ModelAdmin):
    list_display = [
        "serial_number",
        "full_name",
        "email",
        "size",
        "custom_name_display",
        "bulk_order_link",
        "paid_status",
        "coupon_status",
        "created_at",
    ]
    list_filter = [
        "paid",
        "size",
        "bulk_order",
        HasCouponFilter,
    ]
    search_fields = ["full_name", "email", "custom_name"]
    readonly_fields = ["created_at", "updated_at", "serial_number"]
    ordering = ["bulk_order", "serial_number"]
    list_per_page = 20

    def custom_name_display(self, obj):
        """✅ FIX: Only show custom_name if bulk_order has custom_branding_enabled"""
        if obj.bulk_order.custom_branding_enabled and obj.custom_name:
            return obj.custom_name
        return format_html('<span style="color: gray;">-</span>')
    
    custom_name_display.short_description = "Custom Name"

    def paid_status(self, obj):
        if obj.paid:
            return format_html('<span style="color: green; font-weight: bold;">✔ Paid</span>')
        return format_html('<span style="color: red;">✘ Unpaid</span>')

    paid_status.short_description = "Payment Status"

    def coupon_status(self, obj):
        if obj.coupon_used:
            return format_html('<span style="color: blue; font-weight: bold;">🎟️ {}</span>', obj.coupon_used.code)
        return format_html('<span style="color: gray;">-</span>')

    coupon_status.short_description = "Coupon"

    def bulk_order_link(self, obj):
        url = reverse(
            "admin:bulk_orders_bulkorderlink_change", args=[obj.bulk_order.id]
        )
        return format_html('<a href="{}">{}</a>', url, obj.bulk_order.organization_name)

    bulk_order_link.short_description = "Bulk Order"
    
    def get_form(self, request, obj=None, **kwargs):
        """✅ FIX: Filter coupon_used dropdown to show ONLY coupons from the order's bulk_order"""
        form = super().get_form(request, obj, **kwargs)
        
        if obj and 'coupon_used' in form.base_fields:
            # Filter to show only coupons from THIS bulk order
            form.base_fields['coupon_used'].queryset = CouponCode.objects.filter(
                bulk_order=obj.bulk_order,
                is_used=False
            ) | CouponCode.objects.filter(id=obj.coupon_used_id) if obj.coupon_used else CouponCode.objects.filter(
                bulk_order=obj.bulk_order,
                is_used=False
            )
        
        return form


@admin.register(BulkOrderLink)
class BulkOrderLinkAdmin(admin.ModelAdmin):
    list_display = [
        "id",
        "organization_name",
        "slug_display",
        "price_per_item",
        "custom_branding_enabled",
        "payment_deadline",
        "total_orders",
        "total_paid",
        "coupon_count",
        "created_at",
    ]
    list_filter = ["custom_branding_enabled", "created_at", "payment_deadline"]
    search_fields = ["organization_name", "slug"]
    readonly_fields = ["created_at", "updated_at", "slug", "shareable_link"]
    list_per_page = 20
    actions = ['download_pdf_action', 'download_word_action', 'download_excel_action', 'generate_coupons_action']

    inlines = [CouponCodeInline]

    fieldsets = (
        ("Organization Details", {
            "fields": ("organization_name", "slug", "shareable_link", "created_by")
        }),
        ("Order Configuration", {
            "fields": ("price_per_item", "custom_branding_enabled", "payment_deadline")
        }),
        ("Timestamps", {
            "fields": ("created_at", "updated_at"),
            "classes": ("collapse",)
        }),
    )
    
    def changelist_view(self, request, extra_context=None):
        """Store request for use in list_display methods"""
        self._request = request
        return super().changelist_view(request, extra_context)
    
    def change_view(self, request, object_id, form_url='', extra_context=None):
        """Store request for use in readonly_fields methods"""
        self._request = request
        return super().change_view(request, object_id, form_url, extra_context)

    def slug_display(self, obj):
        return format_html('<code style="background: #f0f0f0; padding: 2px 6px; border-radius: 3px;">{}</code>', obj.slug)
    slug_display.short_description = "URL Slug"

    def shareable_link(self, obj):
        """Display the shareable link for easy copying"""
        if obj.slug:
            path = obj.get_shareable_url()
            
            if hasattr(self, '_request'):
                full_url = self._request.build_absolute_uri(path)
            else:
                full_url = path
            
            return format_html(
                '<input type="text" value="{}" readonly style="width: 100%; padding: 5px;" onclick="this.select();" /> '
                '<small style="color: #666;">Click to select and copy</small>',
                full_url
            )
        return "-"
    shareable_link.short_description = "Shareable Link"

    def total_orders(self, obj):
        count = obj.orders.count()
        if count > 0:
            return format_html('<strong>{}</strong>', count)
        return count
    total_orders.short_description = "Total Orders"

    def total_paid(self, obj):
        count = obj.orders.filter(paid=True).count()
        total = obj.orders.count()
        if total > 0:
            percentage = (count / total) * 100
            color = "#2ecc71" if percentage > 80 else "#f39c12" if percentage > 50 else "#e74c3c"
            percentage_str = f"{percentage:.0f}"
            return format_html(
                '<span style="color: {}; font-weight: bold;">{} / {} ({}%)</span>',
                color, count, total, percentage_str
            )
        return "0 / 0"
    total_paid.short_description = "Paid Orders"

    def coupon_count(self, obj):
        total = obj.coupons.count()
        used = obj.coupons.filter(is_used=True).count()
        if total > 0:
            return format_html(
                '<span title="Used / Total">{} / {} <small style="color: #666;">used</small></span>',
                used, total
            )
        return format_html('<span style="color: #999;">No coupons</span>')
    coupon_count.short_description = "Coupons"

    # Admin Actions
    def download_pdf_action(self, request, queryset):
        """Generate PDF for selected bulk orders"""
        if queryset.count() > 1:
            self.message_user(request, "Please select only one bulk order for PDF generation.", messages.WARNING)
            return
        
        bulk_order = queryset.first()
        return self._generate_pdf(request, bulk_order)
    
    download_pdf_action.short_description = "📄 Download PDF"

    def download_word_action(self, request, queryset):
        """Generate Word document for selected bulk orders"""
        if queryset.count() > 1:
            self.message_user(request, "Please select only one bulk order for Word generation.", messages.WARNING)
            return
        
        bulk_order = queryset.first()
        return self._generate_word(request, bulk_order)
    
    download_word_action.short_description = "📝 Download Word"

    def download_excel_action(self, request, queryset):
        """Generate Excel for selected bulk orders"""
        if queryset.count() > 1:
            self.message_user(request, "Please select only one bulk order for Excel generation.", messages.WARNING)
            return
        
        bulk_order = queryset.first()
        return self._generate_excel(request, bulk_order)
    
    download_excel_action.short_description = "📊 Download Excel"

    def generate_coupons_action(self, request, queryset):
        """Generate coupons for selected bulk orders"""
        count = 0
        for bulk_order in queryset:
            if bulk_order.coupons.count() == 0:
                generate_coupon_codes(bulk_order, count=50)
                count += 1
        
        if count > 0:
            self.message_user(
                request,
                f"Successfully generated coupons for {count} bulk order(s).",
                messages.SUCCESS
            )
        else:
            self.message_user(
                request,
                "No bulk orders needed coupon generation (they already have coupons).",
                messages.WARNING
            )
    
    generate_coupons_action.short_description = "🎟️ Generate Coupons (50 per order)"

    # Helper methods for document generation (same as in views.py)
    def _generate_pdf(self, request, bulk_order):
        """Generate PDF - same logic as views.py but for admin"""
        try:
            from weasyprint import HTML
            
            bulk_order = BulkOrderLink.objects.prefetch_related(
                Prefetch(
                    "orders",
                    queryset=OrderEntry.objects.filter(paid=True)
                    .order_by("size", "full_name"),
                )
            ).get(id=bulk_order.id)

            orders = bulk_order.orders.all()
            size_summary = orders.values("size").annotate(count=Count("size")).order_by("size")

            context = {
                "bulk_order": bulk_order,
                "size_summary": size_summary,
                "orders": orders,
                "total_orders": orders.count(),
                "company_name": settings.COMPANY_NAME,
                "company_address": settings.COMPANY_ADDRESS,
                "company_phone": settings.COMPANY_PHONE,
                "company_email": settings.COMPANY_EMAIL,
                "now": timezone.now(),
            }

            html_string = render_to_string("bulk_orders/pdf_template.html", context)
            html = HTML(string=html_string, base_url=request.build_absolute_uri())
            pdf = html.write_pdf()

            response = HttpResponse(pdf, content_type="application/pdf")
            filename = f'bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.pdf'
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            
            logger.info(f"Generated PDF for bulk order: {bulk_order.slug}")
            return response

        except ImportError:
            messages.error(request, "PDF generation not available. Install GTK+ libraries.")
            return None
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            messages.error(request, f"Error generating PDF: {str(e)}")
            return None

    def _generate_word(self, request, bulk_order):
        """Generate Word document - fixed version from views.py"""
        try:
            from docx import Document
            
            bulk_order = BulkOrderLink.objects.prefetch_related(
                Prefetch(
                    "orders",
                    queryset=OrderEntry.objects.filter(paid=True)
                    .order_by("size", "full_name"),
                )
            ).get(id=bulk_order.id)

            doc = Document()
            doc.add_heading(settings.COMPANY_NAME, 0)
            doc.add_heading(f"Bulk Order: {bulk_order.organization_name}", level=1)
            doc.add_paragraph(f"Generated: {timezone.now().strftime('%B %d, %Y - %I:%M %p')}")
            doc.add_paragraph(f"Price per Item: ₦{bulk_order.price_per_item:,.2f}")
            doc.add_paragraph(f"Payment Deadline: {bulk_order.payment_deadline.strftime('%B %d, %Y')}")
            doc.add_paragraph(f"Custom Branding: {'Yes' if bulk_order.custom_branding_enabled else 'No'}")
            doc.add_paragraph("")

            orders = bulk_order.orders.all()

            # Size Summary
            doc.add_heading("Summary by Size", level=2)
            size_summary = orders.values("size").annotate(
                total=Count("id"),
                paid=Count("id", filter=Q(paid=True)),
                coupon=Count("id", filter=Q(coupon_used__isnull=False)),
            ).order_by("size")

            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            header_cells = table.rows[0].cells
            header_cells[0].text = "Size"
            header_cells[1].text = "Total"
            header_cells[2].text = "Paid"
            header_cells[3].text = "Coupon"

            for size_info in size_summary:
                row_cells = table.add_row().cells
                row_cells[0].text = size_info["size"]
                row_cells[1].text = str(size_info["total"])
                row_cells[2].text = str(size_info["paid"])
                row_cells[3].text = str(size_info["coupon"])

            doc.add_paragraph()

            # Orders by Size
            paginator = Paginator(orders, 1000)
            
            for page_num in paginator.page_range:
                page = paginator.page(page_num)
                
                size_groups = {}
                for order in page.object_list:
                    if order.size not in size_groups:
                        size_groups[order.size] = []
                    size_groups[order.size].append(order)

                for size, size_orders in sorted(size_groups.items()):
                    doc.add_heading(f"Size: {size}", level=3)

                    # ✅ FIX: Properly determine column count
                    if bulk_order.custom_branding_enabled:
                        table = doc.add_table(rows=1, cols=4)
                        table.style = "Table Grid"
                        header_cells = table.rows[0].cells
                        header_cells[0].text = "S/N"
                        header_cells[1].text = "Name"
                        header_cells[2].text = "Custom Name"
                        header_cells[3].text = "Status"
                        
                        for idx, order in enumerate(size_orders, 1):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(idx)
                            row_cells[1].text = order.full_name
                            row_cells[2].text = order.custom_name or ''
                            row_cells[3].text = 'Coupon' if order.coupon_used else 'Paid'
                    else:
                        table = doc.add_table(rows=1, cols=3)
                        table.style = "Table Grid"
                        header_cells = table.rows[0].cells
                        header_cells[0].text = "S/N"
                        header_cells[1].text = "Name"
                        header_cells[2].text = "Status"
                        
                        for idx, order in enumerate(size_orders, 1):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(idx)
                            row_cells[1].text = order.full_name
                            row_cells[2].text = 'Coupon' if order.coupon_used else 'Paid'

                    doc.add_paragraph()

            # Footer
            doc.add_paragraph("")
            footer_para = doc.add_paragraph()
            footer_para.add_run(f"{settings.COMPANY_NAME}\n").bold = True
            footer_para.add_run(f"{settings.COMPANY_ADDRESS}\n")
            footer_para.add_run(f"📞 {settings.COMPANY_PHONE} | 📧 {settings.COMPANY_EMAIL}")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            response = HttpResponse(
                buffer.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            filename = f'bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.docx'
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            
            logger.info(f"Generated Word for bulk order: {bulk_order.slug}")
            return response

        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            messages.error(request, f"Error generating Word document: {str(e)}")
            return None

    def _generate_excel(self, request, bulk_order):
        """Generate Excel - complete version with summary at top"""
        try:
            # Optimized query
            bulk_order = BulkOrderLink.objects.prefetch_related(
                Prefetch(
                    "orders",
                    queryset=OrderEntry.objects.filter(paid=True)
                    .order_by("size", "full_name"),
                )
            ).get(id=bulk_order.id)

            output = BytesIO()
            workbook = xlsxwriter.Workbook(output, {"constant_memory": True})

            # ============================================================================
            # FORMATS
            # ============================================================================
            title_format = workbook.add_format({
                'bold': True, 
                'font_size': 16,
                'align': 'left'
            })
            
            subtitle_format = workbook.add_format({
                'bold': True, 
                'font_size': 12,
                'align': 'left'
            })
            
            info_format = workbook.add_format({
                'font_size': 10,
                'align': 'left'
            })
            
            section_header_format = workbook.add_format({
                'bold': True,
                'font_size': 12,
                'bg_color': '#4472C4',
                'font_color': 'white',
                'align': 'center',
                'valign': 'vcenter',
                'border': 1
            })
            
            table_header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#D9E1F2',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })
            
            cell_format = workbook.add_format({
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })
            
            cell_left_format = workbook.add_format({
                'border': 1,
                'align': 'left',
                'valign': 'vcenter'
            })
            
            total_format = workbook.add_format({
                'bold': True,
                'bg_color': '#FFF2CC',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })

            # ============================================================================
            # CREATE WORKSHEET
            # ============================================================================
            worksheet_name = bulk_order.organization_name[:31]
            worksheet = workbook.add_worksheet(worksheet_name)

            # ============================================================================
            # TITLE SECTION
            # ============================================================================
            row = 0
            worksheet.write(row, 0, settings.COMPANY_NAME, title_format)
            row += 1
            
            worksheet.write(row, 0, f"Bulk Order: {bulk_order.organization_name}", subtitle_format)
            row += 1
            
            worksheet.write(row, 0, f"Generated: {timezone.now().strftime('%B %d, %Y - %I:%M %p')}", info_format)
            row += 1
            
            worksheet.write(row, 0, f"Price per Item: ₦{bulk_order.price_per_item:,.2f}", info_format)
            row += 1
            
            worksheet.write(row, 0, f"Payment Deadline: {bulk_order.payment_deadline.strftime('%B %d, %Y')}", info_format)
            row += 1
            
            worksheet.write(row, 0, f"Custom Branding: {'Yes' if bulk_order.custom_branding_enabled else 'No'}", info_format)
            row += 2

            # ============================================================================
            # SIZE SUMMARY SECTION
            # ============================================================================
            orders = bulk_order.orders.all()
            
            size_summary = orders.values("size").annotate(
                total=Count("id"),
                paid=Count("id", filter=Q(paid=True)),
                coupon=Count("id", filter=Q(coupon_used__isnull=False)),
            ).order_by("size")

            # Section header
            worksheet.merge_range(row, 0, row, 3, 'SUMMARY BY SIZE', section_header_format)
            row += 1
            
            # Summary table headers
            summary_headers = ['Size', 'Total', 'Paid', 'Coupon']
            for col, header in enumerate(summary_headers):
                worksheet.write(row, col, header, table_header_format)
            row += 1
            
            # Summary data
            grand_total = 0
            grand_paid = 0
            grand_coupon = 0
            
            for size_data in size_summary:
                worksheet.write(row, 0, size_data['size'], cell_format)
                worksheet.write(row, 1, size_data['total'], cell_format)
                worksheet.write(row, 2, size_data['paid'], cell_format)
                worksheet.write(row, 3, size_data['coupon'], cell_format)
                
                grand_total += size_data['total']
                grand_paid += size_data['paid']
                grand_coupon += size_data['coupon']
                
                row += 1
            
            # Grand total row
            worksheet.write(row, 0, 'TOTAL', total_format)
            worksheet.write(row, 1, grand_total, total_format)
            worksheet.write(row, 2, grand_paid, total_format)
            worksheet.write(row, 3, grand_coupon, total_format)
            row += 2

            # ============================================================================
            # ORDERS SECTION
            # ============================================================================
            
            # Section header
            if bulk_order.custom_branding_enabled:
                worksheet.merge_range(row, 0, row, 4, 'ORDER DETAILS', section_header_format)
            else:
                worksheet.merge_range(row, 0, row, 3, 'ORDER DETAILS', section_header_format)
            row += 1

            # Determine headers
            if bulk_order.custom_branding_enabled:
                headers = ['S/N', 'Size', 'Name', 'Custom Name', 'Status']
            else:
                headers = ['S/N', 'Size', 'Name', 'Status']

            # Write headers
            for col, header in enumerate(headers):
                worksheet.write(row, col, header, table_header_format)
            row += 1

            # Write orders
            serial_number = 1
            for order in orders:
                col = 0
                
                worksheet.write(row, col, serial_number, cell_format)
                col += 1
                
                worksheet.write(row, col, order.size, cell_format)
                col += 1
                
                worksheet.write(row, col, order.full_name, cell_left_format)
                col += 1
                
                if bulk_order.custom_branding_enabled:
                    worksheet.write(row, col, order.custom_name or '', cell_left_format)
                    col += 1
                
                status_text = 'Coupon' if order.coupon_used else 'Paid'
                worksheet.write(row, col, status_text, cell_format)
                
                row += 1
                serial_number += 1

            # ============================================================================
            # COLUMN WIDTHS
            # ============================================================================
            worksheet.set_column(0, 0, 6)   # S/N
            worksheet.set_column(1, 1, 8)   # Size
            worksheet.set_column(2, 2, 30)  # Name
            
            if bulk_order.custom_branding_enabled:
                worksheet.set_column(3, 3, 30)  # Custom Name
                worksheet.set_column(4, 4, 12)  # Status
            else:
                worksheet.set_column(3, 3, 12)  # Status

            # ============================================================================
            # FINALIZE
            # ============================================================================
            workbook.close()
            output.seek(0)

            response = HttpResponse(
                output.read(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            filename = f'bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.xlsx'
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            
            logger.info(f"Generated Excel for bulk order: {bulk_order.slug}")
            return response

        except Exception as e:
            logger.error(f"Error generating Excel: {str(e)}")
            messages.error(request, f"Error generating Excel: {str(e)}")
            return None


@admin.register(CouponCode)
class CouponCodeAdmin(admin.ModelAdmin):
    list_display = ['code', 'bulk_order_link', 'is_used', 'created_at']
    list_filter = ['is_used', 'bulk_order', 'created_at']
    search_fields = ['code', 'bulk_order__organization_name']
    readonly_fields = ['code', 'created_at']
    ordering = ['-created_at']
    
    def bulk_order_link(self, obj):
        url = reverse("admin:bulk_orders_bulkorderlink_change", args=[obj.bulk_order.id])
        return format_html('<a href="{}">{}</a>', url, obj.bulk_order.organization_name)
    bulk_order_link.short_description = "Bulk Order"

    def has_add_permission(self, request):
        return False