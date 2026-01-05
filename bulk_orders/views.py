from rest_framework import viewsets, permissions, filters, status
from rest_framework.decorators import action
from rest_framework.response import Response
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, render
from django.db.models import Count, Q, Prefetch
from django.core.paginator import Paginator
from django.template.loader import render_to_string
from django.conf import settings
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from datetime import datetime, timedelta
import json
import uuid
from .models import BulkOrderLink, OrderEntry, CouponCode
from .serializers import BulkOrderLinkSerializer, OrderEntrySerializer, CouponCodeSerializer
from .utils import generate_receipt, generate_coupon_codes
from payment.utils import initialize_payment, verify_payment
from jmw.background_utils import (
    send_order_confirmation_email,
    send_payment_receipt_email,
    generate_payment_receipt_pdf_task
)
import logging

logger = logging.getLogger(__name__)


class BulkOrderLinkViewSet(viewsets.ModelViewSet):
    serializer_class = BulkOrderLinkSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    queryset = BulkOrderLink.objects.all()
    lookup_field = 'slug'

    def get_queryset(self):
        if self.request.user.is_staff:
             return BulkOrderLink.objects.all()
        if self.request.user.is_authenticated:
            return BulkOrderLink.objects.filter(created_by=self.request.user)
        return BulkOrderLink.objects.none()
    
    @action(detail=True, methods=['post'], permission_classes=[permissions.IsAdminUser])
    def generate_coupons(self, request, slug=None):
        """Generate coupons for a bulk order"""
        bulk_order = self.get_object()
        count = request.data.get('count', 50)
        
        if bulk_order.coupons.count() > 0:
            return Response(
                {"error": f"This bulk order already has {bulk_order.coupons.count()} coupons."},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            coupons = generate_coupon_codes(bulk_order, count=count)
            return Response({
                "message": f"Successfully generated {len(coupons)} coupons",
                "count": len(coupons),
                "sample_codes": [c.code for c in coupons[:5]]
            })
        except Exception as e:
            logger.error(f"Error generating coupons: {str(e)}")
            return Response(
                {"error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAdminUser])
    def download_pdf(self, request, slug=None):
        """Generate PDF summary for this specific bulk order"""
        try:
            from weasyprint import HTML
            
            # Optimized query
            bulk_order = BulkOrderLink.objects.prefetch_related(
                Prefetch(
                    "orders",
                    queryset=OrderEntry.objects.filter(paid=True)
                    .order_by("size", "full_name"),
                )
            ).get(slug=slug)
            
            orders = bulk_order.orders.all()
            
            # Size summary
            size_summary = (
                orders.values("size")
                .annotate(count=Count("size"))
                .order_by("size")
            )
            
            context = {
                'bulk_order': bulk_order,
                'size_summary': size_summary,
                'orders': orders,
                'total_orders': orders.count(),
                'company_name': settings.COMPANY_NAME,
                'company_address': settings.COMPANY_ADDRESS,
                'company_phone': settings.COMPANY_PHONE,
                'company_email': settings.COMPANY_EMAIL,
                'now': timezone.now(),
            }
            
            html_string = render_to_string('bulk_orders/pdf_template.html', context)
            html = HTML(string=html_string)
            pdf = html.write_pdf()
            
            response = HttpResponse(pdf, content_type='application/pdf')
            filename = f'bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.pdf'
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            logger.info(f"Generated PDF for bulk order: {bulk_order.slug}")
            return response
            
        except ImportError:
            return Response(
                {"error": "PDF generation not available. Install GTK+ libraries."},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            return Response(
                {"error": f"Error generating PDF: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAdminUser])
    def download_word(self, request, slug=None):
        """Generate Word document for this specific bulk order"""
        try:
            from docx import Document
            from io import BytesIO
            
            # Optimized query
            bulk_order = BulkOrderLink.objects.prefetch_related(
                Prefetch(
                    "orders",
                    queryset=OrderEntry.objects.filter(paid=True)
                    .order_by("size", "full_name"),
                )
            ).get(slug=slug)
            
            doc = Document()
            
            # Header
            doc.add_heading(settings.COMPANY_NAME, 0)
            doc.add_heading(f'Bulk Order: {bulk_order.organization_name}', level=1)
            doc.add_paragraph(f"Generated: {timezone.now().strftime('%B %d, %Y - %I:%M %p')}")
            doc.add_paragraph(f'Price per Item: ₦{bulk_order.price_per_item:,.2f}')
            doc.add_paragraph(f'Payment Deadline: {bulk_order.payment_deadline.strftime("%B %d, %Y")}')
            doc.add_paragraph(f'Custom Branding: {"Yes" if bulk_order.custom_branding_enabled else "No"}')
            doc.add_paragraph('')
            
            orders = bulk_order.orders.all()
            
            # Size Summary
            doc.add_heading('Summary by Size', level=2)
            size_summary = (
                orders.values("size")
                .annotate(
                    total=Count("id"),
                    paid=Count("id", filter=Q(paid=True)),
                    coupon=Count("id", filter=Q(coupon_used__isnull=False)),
                )
                .order_by("size")
            )
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Size'
            header_cells[1].text = 'Total'
            header_cells[2].text = 'Paid'
            header_cells[3].text = 'Coupon'
            
            for size_info in size_summary:
                row_cells = table.add_row().cells
                row_cells[0].text = size_info['size']
                row_cells[1].text = str(size_info['total'])
                row_cells[2].text = str(size_info['paid'])
                row_cells[3].text = str(size_info['coupon'])
            
            doc.add_paragraph()
            
            # Orders by Size
            paginator = Paginator(orders, 1000)
            
            for page_num in paginator.page_range:
                page = paginator.page(page_num)
                
                size_groups = {}
                for order in page.object_list:
                    if order.size not in size_groups:
                        size_groups[order.size] = []
                    size_groups[order.size].append(order)
                
                for size, size_orders in sorted(size_groups.items()):
                    doc.add_heading(f'Size: {size}', level=3)
                    
                    # ✅ FIX: Properly determine column count based on custom_branding_enabled
                    if bulk_order.custom_branding_enabled:
                        # WITH custom branding: S/N, Name, Custom Name, Status
                        table = doc.add_table(rows=1, cols=4)
                        table.style = 'Table Grid'
                        header_cells = table.rows[0].cells
                        header_cells[0].text = 'S/N'
                        header_cells[1].text = 'Name'
                        header_cells[2].text = 'Custom Name'
                        header_cells[3].text = 'Status'
                        
                        for idx, order in enumerate(size_orders, 1):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(idx)
                            row_cells[1].text = order.full_name
                            row_cells[2].text = order.custom_name or ''
                            row_cells[3].text = 'Coupon' if order.coupon_used else 'Paid'
                    else:
                        # WITHOUT custom branding: S/N, Name, Status
                        table = doc.add_table(rows=1, cols=3)
                        table.style = 'Table Grid'
                        header_cells = table.rows[0].cells
                        header_cells[0].text = 'S/N'
                        header_cells[1].text = 'Name'
                        header_cells[2].text = 'Status'
                        
                        for idx, order in enumerate(size_orders, 1):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(idx)
                            row_cells[1].text = order.full_name
                            row_cells[2].text = 'Coupon' if order.coupon_used else 'Paid'
                    
                    doc.add_paragraph()
            
            # Footer
            doc.add_paragraph("")
            footer_para = doc.add_paragraph()
            footer_para.add_run(f"{settings.COMPANY_NAME}\n").bold = True
            footer_para.add_run(f"{settings.COMPANY_ADDRESS}\n")
            footer_para.add_run(f"📞 {settings.COMPANY_PHONE} | 📧 {settings.COMPANY_EMAIL}")
            
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            response = HttpResponse(
                file_stream.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            filename = f'bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.docx'
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            logger.info(f"Generated Word document for bulk order: {bulk_order.slug}")
            return response
            
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            return Response(
                {"error": f"Error generating Word document: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAdminUser])
    def generate_size_summary(self, request, slug=None):
        """Generate Excel size summary for this specific bulk order"""
        try:
            import xlsxwriter
            from io import BytesIO
            
            # Optimized query
            bulk_order = BulkOrderLink.objects.prefetch_related(
                Prefetch(
                    "orders",
                    queryset=OrderEntry.objects.filter(paid=True)
                    .order_by("size", "full_name"),
                )
            ).get(slug=slug)
            
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output, {'constant_memory': True})
            
            # ============================================================================
            # FORMATS
            # ============================================================================
            title_format = workbook.add_format({
                'bold': True, 
                'font_size': 16,
                'align': 'left'
            })
            
            subtitle_format = workbook.add_format({
                'bold': True, 
                'font_size': 12,
                'align': 'left'
            })
            
            info_format = workbook.add_format({
                'font_size': 10,
                'align': 'left'
            })
            
            section_header_format = workbook.add_format({
                'bold': True,
                'font_size': 12,
                'bg_color': '#4472C4',
                'font_color': 'white',
                'align': 'center',
                'valign': 'vcenter',
                'border': 1
            })
            
            table_header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#D9E1F2',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })
            
            cell_format = workbook.add_format({
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })
            
            cell_left_format = workbook.add_format({
                'border': 1,
                'align': 'left',
                'valign': 'vcenter'
            })
            
            total_format = workbook.add_format({
                'bold': True,
                'bg_color': '#FFF2CC',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })
            
            # ============================================================================
            # CREATE WORKSHEET
            # ============================================================================
            worksheet_name = bulk_order.organization_name[:31]  # Excel limit
            worksheet = workbook.add_worksheet(worksheet_name)
            
            # ============================================================================
            # TITLE SECTION
            # ============================================================================
            row = 0
            worksheet.write(row, 0, settings.COMPANY_NAME, title_format)
            row += 1
            
            worksheet.write(row, 0, f"Bulk Order: {bulk_order.organization_name}", subtitle_format)
            row += 1
            
            worksheet.write(row, 0, f"Generated: {timezone.now().strftime('%B %d, %Y - %I:%M %p')}", info_format)
            row += 1
            
            worksheet.write(row, 0, f"Price per Item: ₦{bulk_order.price_per_item:,.2f}", info_format)
            row += 1
            
            worksheet.write(row, 0, f"Payment Deadline: {bulk_order.payment_deadline.strftime('%B %d, %Y')}", info_format)
            row += 1
            
            worksheet.write(row, 0, f"Custom Branding: {'Yes' if bulk_order.custom_branding_enabled else 'No'}", info_format)
            row += 2  # Extra spacing
            
            # ============================================================================
            # SIZE SUMMARY SECTION (AT TOP, LIKE WORD DOC)
            # ============================================================================
            orders = bulk_order.orders.all()
            
            # Calculate size summary
            size_summary = orders.values("size").annotate(
                total=Count("id"),
                paid=Count("id", filter=Q(paid=True)),
                coupon=Count("id", filter=Q(coupon_used__isnull=False)),
            ).order_by("size")
            
            # Section header
            worksheet.merge_range(row, 0, row, 3, 'SUMMARY BY SIZE', section_header_format)
            row += 1
            
            # Summary table headers
            summary_headers = ['Size', 'Total', 'Paid', 'Coupon']
            for col, header in enumerate(summary_headers):
                worksheet.write(row, col, header, table_header_format)
            row += 1
            
            # Summary data
            grand_total = 0
            grand_paid = 0
            grand_coupon = 0
            
            for size_data in size_summary:
                worksheet.write(row, 0, size_data['size'], cell_format)
                worksheet.write(row, 1, size_data['total'], cell_format)
                worksheet.write(row, 2, size_data['paid'], cell_format)
                worksheet.write(row, 3, size_data['coupon'], cell_format)
                
                grand_total += size_data['total']
                grand_paid += size_data['paid']
                grand_coupon += size_data['coupon']
                
                row += 1
            
            # Grand total row
            worksheet.write(row, 0, 'TOTAL', total_format)
            worksheet.write(row, 1, grand_total, total_format)
            worksheet.write(row, 2, grand_paid, total_format)
            worksheet.write(row, 3, grand_coupon, total_format)
            row += 2  # Extra spacing
            
            # ============================================================================
            # ORDERS SECTION (WITH CONDITIONAL CUSTOM_NAME COLUMN)
            # ============================================================================
            
            # Section header
            if bulk_order.custom_branding_enabled:
                worksheet.merge_range(row, 0, row, 4, 'ORDER DETAILS', section_header_format)
            else:
                worksheet.merge_range(row, 0, row, 3, 'ORDER DETAILS', section_header_format)
            row += 1
            
            # Determine headers based on custom_branding_enabled
            if bulk_order.custom_branding_enabled:
                headers = ['S/N', 'Size', 'Name', 'Custom Name', 'Status']
            else:
                headers = ['S/N', 'Size', 'Name', 'Status']
            
            # Write headers
            for col, header in enumerate(headers):
                worksheet.write(row, col, header, table_header_format)
            row += 1
            
            # Write orders
            start_row = row
            serial_number = 1
            
            for order in orders:
                col = 0
                
                # S/N
                worksheet.write(row, col, serial_number, cell_format)
                col += 1
                
                # Size
                worksheet.write(row, col, order.size, cell_format)
                col += 1
                
                # Name
                worksheet.write(row, col, order.full_name, cell_left_format)
                col += 1
                
                # Custom Name (only if enabled)
                if bulk_order.custom_branding_enabled:
                    worksheet.write(row, col, order.custom_name or '', cell_left_format)
                    col += 1
                
                # Status
                status_text = 'Coupon' if order.coupon_used else 'Paid'
                worksheet.write(row, col, status_text, cell_format)
                
                row += 1
                serial_number += 1
            
            # ============================================================================
            # COLUMN WIDTHS
            # ============================================================================
            worksheet.set_column(0, 0, 6)   # S/N
            worksheet.set_column(1, 1, 8)   # Size
            worksheet.set_column(2, 2, 30)  # Name
            
            if bulk_order.custom_branding_enabled:
                worksheet.set_column(3, 3, 30)  # Custom Name
                worksheet.set_column(4, 4, 12)  # Status
            else:
                worksheet.set_column(3, 3, 12)  # Status
            
            # ============================================================================
            # FINALIZE
            # ============================================================================
            workbook.close()
            output.seek(0)
            
            response = HttpResponse(
                output.read(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            filename = f'bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.xlsx'
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            logger.info(f"Generated Excel for bulk order: {bulk_order.slug}")
            return response
            
        except Exception as e:
            logger.error(f"Error generating Excel: {str(e)}")
            return Response(
                {"error": f"Error generating Excel: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['get'])
    def stats(self, request, slug=None):
        """Get statistics for this specific bulk order"""
        bulk_order = self.get_object()
        total_orders = bulk_order.orders.count()
        paid_orders = bulk_order.orders.filter(paid=True).count()
        
        return Response({
            'organization': bulk_order.organization_name,
            'slug': bulk_order.slug,
            'total_orders': total_orders,
            'paid_orders': paid_orders,
            'unpaid_orders': total_orders - paid_orders,
            'payment_percentage': (paid_orders / total_orders * 100) if total_orders > 0 else 0,
            'total_coupons': bulk_order.coupons.count(),
            'used_coupons': bulk_order.coupons.filter(is_used=True).count(),
            'is_expired': bulk_order.is_expired(),
            'payment_deadline': bulk_order.payment_deadline,
            'custom_branding_enabled': bulk_order.custom_branding_enabled,
        })

    # ✅ NEW: Add orders directly to bulk order (nested route)
    @action(detail=True, methods=['post'], permission_classes=[permissions.AllowAny])
    def submit_order(self, request, slug=None):
        """Submit order for this bulk order (no need for bulk_order_slug in body!)"""
        bulk_order = self.get_object()
        
        # Pass bulk_order via context
        serializer = OrderEntrySerializer(
            data=request.data, 
            context={'bulk_order': bulk_order, 'request': request}
        )
        serializer.is_valid(raise_exception=True)
        order_entry = serializer.save()
        
        return Response(serializer.data, status=status.HTTP_201_CREATED)


class OrderEntryViewSet(viewsets.ModelViewSet):
    """ViewSet for OrderEntry - user's own orders"""
    serializer_class = OrderEntrySerializer
    permission_classes = [permissions.AllowAny]
    
    def get_queryset(self):
        if self.request.user.is_authenticated:
            return OrderEntry.objects.filter(email=self.request.user.email).select_related('bulk_order', 'coupon_used')
        return OrderEntry.objects.none()

    # ✅ Payment initialization endpoint
    @action(detail=True, methods=['post'])
    def initialize_payment(self, request, pk=None):
        """Initialize payment for an OrderEntry"""
        order_entry = self.get_object()
        
        # Check if already paid
        if order_entry.paid:
            return Response(
                {"error": "This order has already been paid"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Generate payment reference
        reference = f"ORDER-{order_entry.bulk_order.id}-{order_entry.id}"
        
        # Calculate amount
        amount = order_entry.bulk_order.price_per_item
        email = order_entry.email
        
        # Build callback URL
        callback_url = request.build_absolute_uri(f"/api/bulk_orders/payment/callback/")
        
        # Initialize payment
        result = initialize_payment(amount, email, reference, callback_url)
        
        if result and result.get('status'):
            return Response({
                "authorization_url": result['data']['authorization_url'],
                "access_code": result['data']['access_code'],
                "reference": reference
            })
        
        return Response(
            {"error": "Payment initialization failed"},
            status=status.HTTP_400_BAD_REQUEST
        )

    @action(detail=True, methods=['get'], permission_classes=[permissions.AllowAny])
    def paid_orders(self, request, slug=None):
        """
        Public page showing all paid orders for social proof.
        Supports HTML view and PDF download.
        """
        bulk_order = self.get_object()
        
        # Get only PAID orders
        paid_orders = bulk_order.orders.filter(paid=True).order_by('-created_at')
        
        # Size summary
        size_summary = (
            paid_orders.values("size")
            .annotate(count=Count("size"))
            .order_by("size")
        )
        
        # Check if download=pdf parameter
        if request.GET.get('download') == 'pdf':
            try:
                from weasyprint import HTML
                
                context = {
                    'bulk_order': bulk_order,
                    'size_summary': size_summary,
                    'paid_orders': paid_orders,
                    'total_paid': paid_orders.count(),
                    'company_name': settings.COMPANY_NAME,
                    'company_address': settings.COMPANY_ADDRESS,
                    'company_phone': settings.COMPANY_PHONE,
                    'company_email': settings.COMPANY_EMAIL,
                    'now': timezone.now(),
                }
                
                html_string = render_to_string('bulk_orders/pdf_template.html', context)
                html = HTML(string=html_string, base_url=request.build_absolute_uri())
                pdf = html.write_pdf()
                
                response = HttpResponse(pdf, content_type='application/pdf')
                filename = f'completed_orders_{bulk_order.slug}.pdf'
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                
                logger.info(f"Generated public paid orders PDF for: {bulk_order.slug}")
                return response
                
            except ImportError:
                return Response(
                    {"error": "PDF generation not available"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
        
        # HTML view
        now = timezone.now()
        days_remaining = (bulk_order.payment_deadline - now).days if bulk_order.payment_deadline > now else 0
        
        context = {
            'bulk_order': bulk_order,
            'size_summary': size_summary,
            'paid_orders': paid_orders,
            'total_paid': paid_orders.count(),
            'recent_orders': paid_orders[:20],  # Show 20 most recent
            'company_name': settings.COMPANY_NAME,
            'now': now,
            'days_remaining': days_remaining,
        }
        
        return render(request, 'bulk_orders/paid_orders_public.html', context)

    @action(detail=True, methods=['get'], permission_classes=[permissions.IsAdminUser])
    def analytics(self, request, slug=None):
        """Admin analytics endpoint for dashboard"""
        bulk_order = self.get_object()
        
        total_orders = bulk_order.orders.count()
        paid_orders = bulk_order.orders.filter(paid=True).count()
        
        # Size breakdown
        size_breakdown = list(
            bulk_order.orders.values('size')
            .annotate(
                total=Count('id'),
                paid=Count('id', filter=Q(paid=True))
            )
            .order_by('size')
        )
        
        # Payment timeline (last 7 days)
        today = timezone.now().date()
        payment_timeline = []
        for i in range(7):
            date = today - timedelta(days=6-i)
            count = bulk_order.orders.filter(
                paid=True,
                updated_at__date=date
            ).count()
            payment_timeline.append({
                'date': date.isoformat(),
                'count': count
            })
        
        # Coupon usage
        total_coupons = bulk_order.coupons.count()
        used_coupons = bulk_order.coupons.filter(is_used=True).count()
        
        return Response({
            'organization': bulk_order.organization_name,
            'slug': bulk_order.slug,
            'overview': {
                'total_orders': total_orders,
                'paid_orders': paid_orders,
                'unpaid_orders': total_orders - paid_orders,
                'payment_percentage': round((paid_orders / total_orders * 100), 2) if total_orders > 0 else 0,
            },
            'size_breakdown': size_breakdown,
            'payment_timeline': payment_timeline,
            'coupons': {
                'total': total_coupons,
                'used': used_coupons,
                'available': total_coupons - used_coupons,
                'usage_percentage': round((used_coupons / total_coupons * 100), 2) if total_coupons > 0 else 0,
            },
            'timeline': {
                'created': bulk_order.created_at,
                'deadline': bulk_order.payment_deadline,
                'is_expired': bulk_order.is_expired(),
                'days_remaining': (bulk_order.payment_deadline - timezone.now()).days if not bulk_order.is_expired() else 0,
            }
        })

class CouponCodeViewSet(viewsets.ModelViewSet):
    serializer_class = CouponCodeSerializer
    permission_classes = [permissions.IsAdminUser]
    queryset = CouponCode.objects.all().select_related('bulk_order')
    
    def get_queryset(self):
        """✅ FIX: Filter coupons by bulk_order if provided"""
        queryset = super().get_queryset()
        bulk_order_slug = self.request.query_params.get('bulk_order_slug')
        
        if bulk_order_slug:
            queryset = queryset.filter(bulk_order__slug=bulk_order_slug)
        
        return queryset
    
    @action(detail=True, methods=['post'])
    def validate_coupon(self, request, pk=None):
        """Validate if a coupon is valid and unused"""
        coupon = self.get_object()
        
        if coupon.is_used:
            return Response({
                "valid": False,
                "message": "This coupon has already been used."
            })
        
        return Response({
            "valid": True,
            "code": coupon.code,
            "bulk_order": coupon.bulk_order.organization_name,
            "bulk_order_slug": coupon.bulk_order.slug,
        })


# ✅ Payment webhook handler for bulk orders
@csrf_exempt
def bulk_order_payment_webhook(request):
    """
    Webhook handler for Paystack payment notifications for bulk orders
    Reference format: ORDER-{bulk_order_id}-{order_entry_id}
    """
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': 'Invalid method'}, status=405)

    try:
        payload = json.loads(request.body)
        logger.info(f"Bulk order payment webhook received: {payload}")

        # Only handle successful charges
        if payload.get('event') != 'charge.success':
            return JsonResponse({'status': 'ignored', 'message': 'Not a charge.success event'})

        data = payload.get('data', {})
        reference = data.get('reference')

        if not reference or not reference.startswith('ORDER-'):
            return JsonResponse({'status': 'error', 'message': 'Invalid reference format'})

        try:
            _, bulk_order_id, order_entry_id = reference.split('-')
        except ValueError:
            logger.error(f"Invalid reference format: {reference}")
            return JsonResponse({'status': 'error', 'message': 'Invalid reference format'})

        # Verify payment with Paystack API
        verification_result = verify_payment(reference)

        if not (
            verification_result
            and verification_result.get('status')
            and verification_result['data']['status'] == 'success'
        ):
            logger.warning(f"Payment verification failed for reference: {reference}")
            return JsonResponse({'status': 'error', 'message': 'Payment verification failed'}, status=400)

        try:
            order_entry = OrderEntry.objects.get(
                id=order_entry_id,
                bulk_order__id=bulk_order_id
            )

            # 🔐 IMPORTANT: Idempotency check
            if order_entry.paid:
                logger.info(f"Webhook already processed for {reference}")
                return JsonResponse({'status': 'success', 'message': 'Already processed'})

            order_entry.paid = True
            order_entry.save(update_fields=['paid'])

            logger.info(
                f"Bulk order payment successful: {reference} "
                f"- OrderEntry {order_entry_id} marked as paid"
            )

            # ✅ SEND PAYMENT RECEIPT EMAIL
            send_payment_receipt_email(order_entry)

            # ✅ GENERATE PDF RECEIPT (ASYNC / BACKGROUND)
            generate_payment_receipt_pdf_task(str(order_entry_id))

            return JsonResponse({
                'status': 'success',
                'message': 'Payment verified and order updated',
                'order_entry_id': str(order_entry_id)
            })

        except OrderEntry.DoesNotExist:
            logger.error(f"OrderEntry not found: {order_entry_id}")
            return JsonResponse({'status': 'error', 'message': 'Order entry not found'}, status=404)

    except json.JSONDecodeError:
        logger.error("Invalid JSON in webhook payload")
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)
    except Exception as e:
        logger.exception("Error processing bulk order payment webhook")
        return JsonResponse({'status': 'error', 'message': 'Internal server error'}, status=500)
